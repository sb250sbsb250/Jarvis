"""
Word 工具（原子工具版）

原子工具:
  word_read   — 读取 Word 文档
  word_write  — 写入 Word 文档
"""

import os
import logging
from typing import List

from engine.tool.base import (
    BaseTool, ToolDefinition, ToolParameter, ToolResult,
    CATEGORY_DATA,
)

logger = logging.getLogger(__name__)


class WordTool(BaseTool):
    """Word 文档工具集"""

    def __init__(self):
        self._handlers = {
            "word_read": self._handle_read,
            "word_write": self._handle_write,
        }
        for t in self.tools:
            t.handler = self._handlers.get(t.name)

    @property
    def name(self) -> str:
        return "word"

    @property
    def category(self) -> str:
        return CATEGORY_DATA

    @property
    def tools(self) -> List[ToolDefinition]:
        return [
            ToolDefinition(
                name="word_read",
                description="""读取 Word 文档（.docx）的文本内容。
自动提取所有段落的文本，忽略格式信息。

使用场景：
- 读取 .docx 报告的文本内容
- 从 Word 文档中提取信息

不适用场景：
- 读取 .doc 格式（旧版 Word）→ 建议先转换为 .docx
- 读取 PDF → 用 pdf_read""",
                parameters=[
                    ToolParameter("path", "string", "Word 文件路径（仅支持 .docx 格式）", required=True),
                ],
                is_read=True,
                examples=[
                    'word_read(path="report.docx")',
                ],
                constraints=[
                    "仅支持 .docx 格式，不支持 .doc（旧版 Word）",
                    "只提取纯文本内容，不包含图片/表格/样式",
                    "需要安装 python-docx：pip install python-docx",
                ],
            ),
            ToolDefinition(
                name="word_write",
                description="""写入新的 Word 文档（.docx）。
将文本内容写入一个新的 Word 文档，每行文本生成一个段落。

使用场景：
- 生成 Word 格式的报告
- 创建简单的 .docx 文档

不适用场景：
- 修改已有 Word 文档 → 用 word_read + 修改后 word_write 覆盖""",
                parameters=[
                    ToolParameter("path", "string", "输出文件路径（必须 .docx 结尾）", required=True),
                    ToolParameter("text", "string", "文档内容，换行符 \\n 表示段落分隔", required=True),
                ],
                examples=[
                    'word_write(path="output.docx", text="标题\\n\\n正文内容\\n\\n结尾")',
                ],
                constraints=[
                    "每行文本（被 \\n 分隔）会生成一个独立段落",
                    "不支持图片/表格/复杂格式，纯文本输出",
                    "如果文件已存在会覆盖",
                    "需要安装 python-docx：pip install python-docx",
                ],
            ),
        ]

    async def execute(self, call_id: str, tool_name: str, **kwargs) -> ToolResult:
        handler = self._handlers.get(tool_name)
        if not handler:
            return ToolResult.fail(call_id, tool_name, f"未知工具: {tool_name}")
        try:
            path = kwargs.get("path", "")
            if tool_name == "word_read" and (not path or not os.path.exists(path)):
                return ToolResult.fail(call_id, tool_name, f"文件不存在: {path}")
            return await handler(call_id, **kwargs)
        except Exception as e:
            return ToolResult.fail(call_id, tool_name, str(e))

    async def _handle_read(self, call_id: str, path: str) -> ToolResult:
        try:
            from docx import Document
            doc = Document(path)
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return ToolResult.ok(call_id, "word_read", {
                "path": path,
                "content": text[:15000],
                "total_paragraphs": len(doc.paragraphs),
            })
        except ImportError:
            return ToolResult.fail(call_id, "word_read", "需要 python-docx: pip install python-docx")
        except Exception as e:
            return ToolResult.fail(call_id, "word_read", str(e))

    async def _handle_write(self, call_id: str, path: str, text: str) -> ToolResult:
        try:
            from docx import Document
            doc = Document()
            for paragraph in text.split("\n"):
                doc.add_paragraph(paragraph)
            doc.save(path)
            return ToolResult.ok(call_id, "word_write", {
                "path": path,
                "size": len(text),
                "status": "已写入",
            })
        except ImportError:
            return ToolResult.fail(call_id, "word_write", "需要 python-docx: pip install python-docx")
        except Exception as e:
            return ToolResult.fail(call_id, "word_write", str(e))
